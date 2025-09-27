# generate_test_plan.py
import os
import json
import re
import time
import argparse
import config # Import API keys
import pandas as pd
from docx import Document
from docx.shared import Inches
from PIL import Image
from dotenv import load_dotenv

# --- Third-Party Libraries ---
import google.generativeai as genai
from google.api_core import exceptions
from tavily import TavilyClient
from langchain_community.vectorstores import Chroma
from langchain_google_genai import GoogleGenerativeAIEmbeddings

# --- Environment and API Configuration ---
os.environ["GOOGLE_API_KEY"] = config.GOOGLE_API_KEY
os.environ["TAVILY_API_KEY"] = config.TAVILY_API_KEY
genai.configure(api_key=config.GOOGLE_API_KEY)

load_dotenv()  # Load variables from .env into the environment

# --- Constants ---
CHROMA_PATH = "chroma_db"
GENERATION_CONFIG = {
    "temperature": 0.2,
    "top_p": 0.9,
    "top_k": 32,
}

# --- Helper Functions ---
def read_file_content(file_path):
    """Reads the content of a given file, handling CSVs with pandas."""
    if not file_path or not os.path.exists(file_path):
        print(f"⚠️ Warning: File not found or path is empty: {file_path}")
        return ""
    try:
        if file_path.lower().endswith('.csv'):
            return pd.read_csv(file_path, on_bad_lines='skip').to_string()
        else:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
    except Exception as e:
        print(f"❌ Error reading {file_path}: {e}")
        return ""

def call_gemini_with_retry(model, prompt, retries=3, base_delay=15):
    """Calls the Gemini API with exponential backoff for rate limit errors."""
    for i in range(retries):
        try:
            response = model.generate_content(prompt, generation_config=GENERATION_CONFIG)
            return response.text
        except exceptions.ResourceExhausted as e:
            wait_time = base_delay * (2 ** i)
            print(f"Rate limit hit. Retrying in {wait_time} seconds... ({i+1}/{retries})")
            time.sleep(wait_time)
        except Exception as e:
            print(f"An unexpected error occurred: {e}")
            return None
    print("❌ All retries failed. Could not get a response from the API.")
    return None

# --- RAG and Agent Functions ---
def get_relevant_context_from_rag(query, k=8):
    """Retrieves relevant document chunks from the Chroma vector store."""
    print("🧠 Querying RAG vector store for relevant context...")
    if not os.path.exists(CHROMA_PATH):
        return "⚠️ RAG database not found. Please run ingest_docs.py first."
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    db = Chroma(persist_directory=CHROMA_PATH, embedding_function=embeddings)
    results = db.similarity_search(query, k=k)
    context = "\n---\n".join([doc.page_content for doc in results])
    return context

def run_batch_datasheet_agent(part_numbers):
    """Searches for datasheets and extracts key specs for a list of components."""
    print(f"🤖 Running Datasheet Agent for: {part_numbers}...")
    tavily_client = TavilyClient(api_key=os.environ["TAVILY_API_KEY"])
    model = genai.GenerativeModel('gemini-1.5-pro-latest')

    all_context = ""
    for part in part_numbers:
        try:
            print(f"   -> Searching online for {part} datasheet...")
            search_results = tavily_client.search(
                query=f'datasheet specifications for "{part}" recommended operating voltage and absolute maximum ratings',
                search_depth="advanced"
            )
            context = "\n".join([result['content'] for result in search_results['results']])
            all_context += f"\n\n--- Search Results for {part} ---\n{context}"
        except Exception as e:
            print(f"   -> ⚠️ Could not perform web search for {part}: {e}")

    prompt = f"""
    Based on the following datasheet search results, extract the 'recommended operating voltage range' and the 'absolute maximum VCC/VDD voltage' for each component.
    Provide the answer as a single, clean JSON object where each key is a part number. Be precise.

    SEARCH RESULTS:
    {all_context}

    JSON OUTPUT:
    """
    specs_json_str = call_gemini_with_retry(model, prompt)
    try:
        return json.loads(specs_json_str.strip().replace('```json', '').replace('```', ''))
    except (json.JSONDecodeError, AttributeError):
        print("❌ Agent failed to return valid JSON for datasheets.")
        return {}

def generate_annotated_image(component_ref, pcb_image_path):
    """Asks Gemini 1.5 Pro to highlight a component on a PCB image."""
    print(f"🎨 Generating annotated image for '{component_ref}'...")
    try:
        model = genai.GenerativeModel('gemini-1.5-pro-latest')
        pcb_image = Image.open(pcb_image_path)
        prompt = [
            f"On this PCB image, please draw a single, bright red circle to clearly highlight the component with the reference designator '{component_ref}'.",
            pcb_image
        ]
        response = model.generate_content(prompt, stream=False)
        output_filename = f"annotated_{component_ref}.png"
        with open(output_filename, 'wb') as f:
            f.write(response.parts[0].inline_data.data)
        print(f"   -> Saved annotated image to {output_filename}")
        return output_filename
    except Exception as e:
        print(f"   -> ❌ Error generating image for {component_ref}: {e}")
        return None

# --- Main Workflow ---
def generate_test_plan(netlist_path, bom_path, layout_image_path, output_path):
    """Main function to generate the PCB test plan."""
    model = genai.GenerativeModel('gemini-1.5-pro-latest')

    # 1. Read raw input files
    netlist_content = read_file_content(netlist_path)
    bom_content = read_file_content(bom_path)

    # 2. Get context from RAG
    rag_query = "Extract all requirements for power, microcontrollers, communication buses (SPI, I2C, UART), connectors, and key ICs."
    rag_context = get_relevant_context_from_rag(rag_query)

    # 3. Stage 1: Analysis and JSON Extraction
    print("--- STAGE 1: Analyzing Design and Extracting Key Info ---")
    stage1_prompt = f"""
    You are an expert electronics test engineer. Analyze the provided RAG context, netlist, and BOM.
    Extract key information into a comprehensive JSON object with keys: power_rails, ics, communication_buses, connectors, user_io, clocks.
    For the 'ics' list, ensure you include the reference designator and the full part number.

    **RAG Context from Design Documents:**
    {rag_context}
    ---
    **Netlist Data:**
    {netlist_content}
    ---
    **BOM Data:**
    {bom_content}
    """
    json_summary_str = call_gemini_with_retry(model, stage1_prompt)
    if not json_summary_str: return

    # 4. Enrich with Datasheet Agent
    print("\n--- STAGE 1.5: Enriching Data with Datasheet Agent ---")
    try:
        summary_data = json.loads(json_summary_str.strip().replace('```json', '').replace('```', ''))
        part_numbers_to_check = [ic['part_number'] for ic in summary_data.get('ics', []) if 'part_number' in ic]
        if part_numbers_to_check:
            datasheet_specs = run_batch_datasheet_agent(part_numbers_to_check)
            for ic in summary_data.get('ics', []):
                if ic.get('part_number') in datasheet_specs:
                    ic['voltage_specs'] = datasheet_specs[ic['part_number']]
        enriched_json_summary = json.dumps(summary_data, indent=2)
    except json.JSONDecodeError:
        print("❌ Could not parse Stage 1 JSON. Proceeding without enrichment.")
        enriched_json_summary = json_summary_str

    # 5. Stage 2: Test Plan Generation
    print("\n--- STAGE 2: Generating Final Test Plan ---")
    stage2_prompt = f"""
    You are a senior test procedure writer creating a Bring-Up Test Procedure for a junior engineer.
    Based on the following JSON summary, create a detailed, step-by-step test plan in Markdown format.
    The plan must include: Document Header, Required Equipment, Setup Instructions, and logical Test Sections.
    For each section, create a table with columns: "Step", "Action", "Measurement Point(s)", "Expected Result", and "Pass/Fail".
    Use the 'voltage_specs' from the JSON to create precise pass/fail criteria.
    Prioritize critical paths: power rails first, then clocks, then microcontroller, then peripherals.

    **PCB Design Summary:**
    {enriched_json_summary}
    """
    markdown_plan = call_gemini_with_retry(model, stage2_prompt)
    if not markdown_plan: return

    # 6. Save to Word Document with Annotated Images
    print("\n--- STAGE 3: Formatting Output and Generating Images ---")
    save_to_word_with_images(markdown_plan, output_path, layout_image_path)

def save_to_word_with_images(markdown_text, filename, pcb_image_path):
    """Saves the markdown plan to a .docx file, adding annotated images where possible."""
    doc = Document()
    doc.add_heading('PCB Bring-Up Test Procedure', 0)
    ref_des_pattern = re.compile(r'\b([UQRCDJTP][0-9]+)\b')
    processed_components = set()

    for line in markdown_text.split('\n'):
        doc.add_paragraph(line)
        # Check if the line contains a component and we have an image to work with
        if pcb_image_path and '|' in line: # Process components mentioned in table rows
            found_components = ref_des_pattern.findall(line)
            if found_components:
                component_to_find = found_components[0]
                if component_to_find not in processed_components:
                    image_file = generate_annotated_image(component_to_find, pcb_image_path)
                    if image_file:
                        doc.add_picture(image_file, width=Inches(4.0))
                        processed_components.add(component_to_find)

    doc.save(filename)
    print(f"\n✅ Success! Test plan saved to '{filename}'")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="AI-Powered PCB Test Plan Generator")
    parser.add_argument("--netlist", required=True, help="Path to the netlist file.")
    parser.add_argument("--bom", help="Path to the Bill of Materials CSV file.")
    parser.add_argument("--layout_image", help="Path to the PCB layout image (.png, .jpg) for annotation.")
    parser.add_argument("--output", default="Generated_Test_Plan.docx", help="Output filename for the Word document.")
    args = parser.parse_args()

    generate_test_plan(args.netlist, args.bom, args.layout_image, args.output)